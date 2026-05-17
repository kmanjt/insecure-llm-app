"""Document ingest pipeline.

Order of operations for every upload:

1. Best-effort text extraction (binary formats parsed in-process: pypdf
   for PDFs, python-docx for .docx, openpyxl for .xlsx). Plain text is
   decoded directly.
2. If the firewall is enabled and we have any extracted text, send it
   to SonnyLabs at the ``document`` surface. A positive detection
   short-circuits the rest of the pipeline — the file never lands in
   blob storage, the vector store, or AI Search.
3. Upload to Azure Blob Storage (central source of truth).
4. Push to the Foundry agent's vector store (for ``file_search``).
5. Optionally index plaintext content into AI Search (for the agent's
   ``azure_ai_search`` tool).
"""
from io import BytesIO

from . import firewall
from .blob_client import upload_blob
from .foundry_client import upload_to_vector_store
from .search_client import index_text_document

_TEXTUAL_EXTS = {".txt", ".md", ".csv", ".log", ".json", ".yaml", ".yml"}
_SCAN_PREVIEW_CHARS = 50_000
_PDF_MAX_PAGES = 20
_XLSX_MAX_SHEETS = 5
_XLSX_MAX_ROWS = 500


def _is_textual(filename: str) -> bool:
    name = filename.lower()
    return any(name.endswith(ext) for ext in _TEXTUAL_EXTS)


def _extract_text_for_scan(filename: str, content: bytes) -> str | None:
    """Pull text out of common document formats so the firewall has
    something to scan. Returns ``None`` when extraction isn't possible
    (unsupported type, e.g. image) or fails."""
    name = filename.lower()

    if _is_textual(name):
        return content[: _SCAN_PREVIEW_CHARS * 2].decode("utf-8", errors="ignore")

    try:
        if name.endswith(".pdf"):
            from pypdf import PdfReader
            reader = PdfReader(BytesIO(content))
            parts: list[str] = []
            for page in reader.pages[:_PDF_MAX_PAGES]:
                try:
                    parts.append(page.extract_text() or "")
                except Exception:  # noqa: BLE001 — corrupt page, skip
                    pass
            return "\n".join(parts)[:_SCAN_PREVIEW_CHARS]

        if name.endswith(".docx"):
            from docx import Document
            doc = Document(BytesIO(content))
            text = "\n".join(p.text for p in doc.paragraphs)
            # Tables too — easy win for spreadsheety .docx docs.
            for table in doc.tables:
                for row in table.rows:
                    text += "\n" + " | ".join(cell.text for cell in row.cells)
            return text[:_SCAN_PREVIEW_CHARS]

        if name.endswith(".xlsx"):
            from openpyxl import load_workbook
            wb = load_workbook(BytesIO(content), read_only=True, data_only=True)
            parts = []
            for sheet in wb.worksheets[:_XLSX_MAX_SHEETS]:
                for row in sheet.iter_rows(max_row=_XLSX_MAX_ROWS, values_only=True):
                    parts.append(" ".join(str(c) for c in row if c is not None))
                if len("\n".join(parts)) > _SCAN_PREVIEW_CHARS:
                    break
            return "\n".join(parts)[:_SCAN_PREVIEW_CHARS]
    except Exception as exc:  # noqa: BLE001 — any extraction problem is non-fatal
        print(f"[ingest] text extraction failed for {filename}: {exc}", flush=True)
        return None

    return None  # images, .doc, .xls, anything else


def ingest_document(filename: str, content: bytes) -> dict:
    # 1+2. Extract + scan (best-effort; binary types with no extractor
    # available, like images, pass through unscanned with a note in the
    # response).
    firewall_state: dict | None = None
    if firewall.is_enabled():
        extracted = _extract_text_for_scan(filename, content)
        if extracted and extracted.strip():
            try:
                tag = firewall.check_or_raise("document", extracted)
                firewall_state = {"scanned": True, "scan_id": tag, "extracted_chars": len(extracted)}
            except firewall.FirewallBlock as fb:
                return {
                    "filename": filename,
                    "blocked": True,
                    "block_surface": fb.surface,
                    "block_reason": fb.summary,
                    "scan_id": fb.scan_id,
                }
        elif extracted is None:
            firewall_state = {"scanned": False, "reason": "no extractor for this file type"}
            print(f"[ingest] no extractor for {filename}; uploading unscanned", flush=True)
        else:
            firewall_state = {"scanned": False, "reason": "extracted text was empty"}

    # 3. Central blob (source of truth).
    blob_url = upload_blob(filename, content)

    # 4. Foundry vector store (file_search tool). Some file types (e.g.
    # images) aren't supported by the agent's file_search; rather than
    # 500-ing the whole upload we capture the error and continue — the
    # raw file is still in blob storage and can be referenced by the
    # agent via the azure_ai_search tool if it was indexed.
    vector_store_file_id = None
    vector_store_error = None
    try:
        vector_store_file_id = upload_to_vector_store(filename, content)
    except Exception as exc:  # noqa: BLE001
        vector_store_error = f"{type(exc).__name__}: {exc}"
        print(f"[ingest] vector store upload failed for {filename}: {vector_store_error}", flush=True)

    # 5. AI Search (plaintext only — keeps the index simple).
    indexed_in_search = False
    if _is_textual(filename):
        text = content.decode("utf-8", errors="ignore")
        index_text_document(filename, text)
        indexed_in_search = True

    out = {
        "filename": filename,
        "blob_url": blob_url,
        "vector_store_file_id": vector_store_file_id,
        "indexed_in_search": indexed_in_search,
    }
    if vector_store_error is not None:
        out["vector_store_error"] = vector_store_error
    if firewall_state is not None:
        out["firewall"] = firewall_state
    return out
